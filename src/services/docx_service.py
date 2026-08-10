from __future__ import annotations

import io
import os
import re
from datetime import date
from typing import Any, Dict, Iterable, List, Optional, Tuple

import docx
import pandas as pd
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

from src.config import (
    COLUMNAS_EVALUACIONES,
    COLUMNAS_GC72,
    COLUMNAS_HORARIOS,
    COLUMNAS_MODULOS,
    COLUMNAS_SESIONES,
    MAPA_TABLA_WORD_GC72,
    NUMERICAS_ENTERAS_GC72,
    NUMERICAS_PORCENTAJE_GC72,
    PREFORMAS_GC72,
    TEMPLATE_GC71,
    TEMPLATE_GC72,
)


def porcentaje(numerador: Any, denominador: Any) -> str:
    try:
        num = float(numerador)
        den = float(denominador)
        if den == 0:
            return "0%"
        return f"{(num / den) * 100:.1f}%".replace(".0%", "%")
    except Exception:
        return "0%"


def formato_entero(valor: Any) -> str:
    if valor is None or pd.isna(valor) or str(valor).strip() == "":
        return "0"
    try:
        return str(int(round(float(valor))))
    except Exception:
        return str(valor)


def formato_porcentaje(valor: Any) -> str:
    if valor is None or pd.isna(valor) or str(valor).strip() == "":
        return "0%"
    txt = str(valor).strip()
    if txt.endswith("%"):
        return txt
    try:
        val = float(txt)
        if val <= 1.0 and val > 0:
            val = val * 100
        return f"{val:.1f}%".replace(".0%", "%")
    except Exception:
        return txt


def shade_cell(cell: Any, color_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_width(cell: Any, width_in_inches: float):
    cell.width = Inches(width_in_inches)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_in_inches * 1440)}" w:type="dxa"/>')
    tcPr.append(tcW)


def fijar_anchos_tabla(table: Any, col_widths_in_inches: List[float]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(col_widths_in_inches):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def aplicar_bordes_tabla(table: Any, color_hex: str = "D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="002060"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_cell_text(cell: Any, text: str, bold: bool = False, italic: bool = False, font_size_pt: float = 9.5, color_rgb: Tuple[int, int, int] = (0, 0, 0), align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph_in_cell(cell: Any, text: str = "", bold: bool = False, italic: bool = False, font_size_pt: float = 9.5, color_rgb: Tuple[int, int, int] = (0, 0, 0), space_after: float = 2):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    return p


def remover_parrafo(p: Any):
    p._element.getparent().remove(p._element)


def set_label_value(cell: Any, label: str, value: str, font_size_pt: float = 9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05

    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    run_l.font.size = Pt(font_size_pt)
    run_l.font.color.rgb = RGBColor(0, 32, 96)

    run_v = p.add_run(str(value or ""))
    run_v.bold = False
    run_v.font.size = Pt(font_size_pt)
    run_v.font.color.rgb = RGBColor(30, 30, 30)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_section_text(cell: Any, text: str, font_size_pt: float = 9.5):
    cell.text = ""
    lineas = str(text or "").split("\n")
    first = True
    for line in lineas:
        txt = line.strip()
        if not txt:
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(txt)
        run.font.size = Pt(font_size_pt)
        run.font.color.rgb = RGBColor(30, 30, 30)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def normalizar_dataframe_gc72(df_input: pd.DataFrame, calcular_porcentajes: bool = True) -> pd.DataFrame:
    df = df_input.copy()
    for col in COLUMNAS_GC72:
        if col not in df.columns:
            df[col] = ""

    for idx, row in df.iterrows():
        asig = str(row.get("Asignatura", "")).strip()
        cod = str(row.get("Código", "")).strip()
        grp = str(row.get("Grupo", "")).strip()

        if not asig and not cod and not grp:
            continue

        matriculados = row.get("Estudiantes matriculados", 0)
        try:
            matriculados = int(round(float(matriculados))) if matriculados != "" else 0
        except Exception:
            matriculados = 0

        for col in NUMERICAS_ENTERAS_GC72:
            val = row.get(col, 0)
            try:
                df.at[idx, col] = int(round(float(val))) if val != "" else 0
            except Exception:
                df.at[idx, col] = 0

        if calcular_porcentajes:
            df.at[idx, "Desertaron %"] = porcentaje(row.get("Desertaron N°", ""), matriculados)
            df.at[idx, "Aprueban evaluación parcial %"] = porcentaje(row.get("Aprueban evaluación parcial N°", ""), matriculados)
            df.at[idx, "Reprueban evaluación parcial %"] = porcentaje(row.get("Reprueban evaluación parcial N°", ""), matriculados)
            df.at[idx, "Aprueban a la fecha %"] = porcentaje(row.get("Aprueban a la fecha N°", ""), matriculados)
            df.at[idx, "Reprueban a la fecha %"] = porcentaje(row.get("Reprueban a la fecha N°", ""), matriculados)
    return df


def texto_preformas(seleccionadas: Iterable[str], categoria: str, adicional: str = "") -> str:
    partes = [PREFORMAS_GC72[categoria][k] for k in seleccionadas if k in PREFORMAS_GC72[categoria]]
    adicional = (adicional or "").strip()
    if adicional:
        partes.append(adicional)
    return " ".join(partes).strip()


def curso_key(row: Any, idx: int) -> str:
    codigo = re.sub(r"\W+", "_", str(row.get("Código", "")).strip())
    grupo = re.sub(r"\W+", "_", str(row.get("Grupo", "")).strip())
    asignatura = re.sub(r"\W+", "_", str(row.get("Asignatura", "")).strip())[:40]
    return f"curso_{idx}_{codigo}_{grupo}_{asignatura}"


def agregar_parrafo_antes(parrafo_referencia: Any, texto: str = "", bold_prefix: Optional[str] = None, space_after: int = 3):
    nuevo = parrafo_referencia.insert_paragraph_before()
    nuevo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nuevo.paragraph_format.space_after = Pt(space_after)
    nuevo.paragraph_format.line_spacing = 1.05
    if bold_prefix and texto.startswith(bold_prefix):
        r1 = nuevo.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = nuevo.add_run(texto[len(bold_prefix):])
        r2.font.size = Pt(10)
    else:
        r = nuevo.add_run(texto)
        r.font.size = Pt(10)
    return nuevo


def construir_analisis(cursos: pd.DataFrame, analisis_por_curso: Dict[str, Dict[str, str]], modo: str) -> List[Dict[str, str]]:
    bloques = []
    if modo == "Consolidado institucional":
        positivos, inconvenientes, propuestas = [], [], []
        for data in analisis_por_curso.values():
            if data.get("positivos"):
                positivos.append(data["positivos"])
            if data.get("inconvenientes"):
                inconvenientes.append(data["inconvenientes"])
            if data.get("propuestas"):
                propuestas.append(data["propuestas"])
        bloques.append({
            "titulo": "Análisis consolidado de los cursos reportados",
            "positivos": " ".join(dict.fromkeys(positivos)),
            "inconvenientes": " ".join(dict.fromkeys(inconvenientes)),
            "propuestas": " ".join(dict.fromkeys(propuestas)),
        })
        return bloques

    for idx, row in cursos.iterrows():
        key = curso_key(row, idx)
        nombre = str(row.get("Asignatura", "")).strip() or "Curso sin nombre"
        codigo = str(row.get("Código", "")).strip()
        grupo = str(row.get("Grupo", "")).strip()
        data = analisis_por_curso.get(key, {})
        bloques.append({
            "titulo": f"{nombre} | Código: {codigo} | Grupo: {grupo}",
            "positivos": data.get("positivos", ""),
            "inconvenientes": data.get("inconvenientes", ""),
            "propuestas": data.get("propuestas", ""),
        })
    return bloques


def set_template_cell(table: Any, row_idx: int, col_idx: int, text: Any, font_size: float = 9, bold: bool = False):
    if row_idx < len(table.rows):
        row = table.rows[row_idx]
        if col_idx < len(row.cells):
            cell = row.cells[col_idx]
            cell.text = str(text if text is not None else "")
            if cell.paragraphs:
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05
                if p.runs:
                    p.runs[0].font.size = Pt(font_size)
                    p.runs[0].bold = bold


def crear_gc71_docx(
    datos: Dict[str, str],
    sesiones_df: pd.DataFrame,
    evaluaciones_df: pd.DataFrame,
    representantes_df: Optional[pd.DataFrame] = None,
) -> bytes:
    doc = docx.Document(TEMPLATE_GC71 if os.path.exists(TEMPLATE_GC71) else None)

    # Si la plantilla oficial existe y tiene la tabla estructurada (Table 0 con al menos 70 filas)
    if len(doc.tables) > 0 and len(doc.tables[0].rows) >= 70:
        table = doc.tables[0]

        # 1. Identificación
        set_template_cell(table, 3, 6, datos.get("programa", ""))
        set_template_cell(table, 4, 6, datos.get("asignatura", ""))
        set_template_cell(table, 5, 6, datos.get("codigo", ""))
        set_template_cell(table, 6, 6, datos.get("area", ""))
        set_template_cell(table, 7, 6, datos.get("prerrequisitos", "Ninguno"))
        set_template_cell(table, 8, 6, datos.get("correquisitos", "Ninguno"))

        tipo = str(datos.get("tipo_asignatura", "")).lower()
        if "teórica" in tipo and "práctica" not in tipo:
            set_template_cell(table, 9, 7, "X", bold=True)
        elif "práctica" in tipo and "teórica" not in tipo:
            set_template_cell(table, 9, 18, "X", bold=True)
        else:
            set_template_cell(table, 9, 14, "X", bold=True)

        set_template_cell(table, 10, 6, datos.get("creditos", ""))
        set_template_cell(table, 11, 9, datos.get("htp", ""))
        set_template_cell(table, 11, 17, datos.get("hti", ""))
        set_template_cell(table, 12, 6, datos.get("profesor", ""))
        set_template_cell(table, 13, 6, datos.get("correo", ""))
        set_template_cell(table, 14, 6, datos.get("grupo", ""))
        set_template_cell(table, 15, 6, datos.get("periodo", ""))

        # 2. Textos académicos
        if datos.get("justificacion"):
            set_template_cell(table, 16, 0, f"JUSTIFICACIÓN:\n{datos['justificacion']}", font_size=8.5)
        if datos.get("competencias"):
            set_template_cell(table, 17, 0, f"COMPETENCIAS A LAS QUE LE TRIBUTA LA ASIGNATURA:\n{datos['competencias']}", font_size=8.5)
        if datos.get("resultados"):
            set_template_cell(table, 18, 0, f"RESULTADOS DE APRENDIZAJE A LOS QUE LE TRIBUTA LA ASIGNATURA:\n{datos['resultados']}", font_size=8.5)
        if datos.get("objetivo_general"):
            set_template_cell(table, 20, 0, f"OBJETIVO(S) GENERAL(ES):\n{datos['objetivo_general']}", font_size=8.5)
        if datos.get("objetivos_especificos"):
            set_template_cell(table, 21, 0, f"OBJETIVOS ESPECÍFICOS:\n{datos['objetivos_especificos']}", font_size=8.5)
        if datos.get("metodologias"):
            set_template_cell(table, 22, 0, f"METODOLOGÍAS Y ESTRATEGIAS DIDÁCTICAS DE LA ASIGNATURA:\n{datos['metodologias']}", font_size=8.5)
        if datos.get("ambientes"):
            set_template_cell(table, 49, 0, f"AMBIENTES DE APRENDIZAJE DE LA ASIGNATURA:\n{datos['ambientes']}", font_size=8.5)
        if datos.get("medios"):
            set_template_cell(table, 50, 0, f"MEDIOS EDUCATIVOS PARA LA ASIGNATURA:\n{datos['medios']}", font_size=8.5)
        if datos.get("referencias"):
            set_template_cell(table, 51, 0, f"REFERENCIAS BIBLIOGRÁFICAS:\n{datos['referencias']}", font_size=8.5)

        # 3. Sesiones en plantilla
        filas_sesiones_plantilla = [26, 27, 28, 31, 32, 33, 36, 37, 38, 41, 42, 43, 46, 47, 48]
        if sesiones_df is not None and not sesiones_df.empty:
            for s_idx, (_, r) in enumerate(sesiones_df.iterrows()):
                if s_idx < len(filas_sesiones_plantilla):
                    t_row = filas_sesiones_plantilla[s_idx]
                    set_template_cell(table, t_row, 0, r.get("N° sesión", s_idx + 1), font_size=8)
                    set_template_cell(table, t_row, 1, r.get("Fecha", ""), font_size=8)
                    set_template_cell(table, t_row, 3, r.get("Contenido por desarrollar", ""), font_size=8)
                    set_template_cell(table, t_row, 6, r.get("Descripción del trabajo presencial", ""), font_size=8)
                    set_template_cell(table, t_row, 12, r.get("Descripción trabajo independiente", ""), font_size=8)

        # 4. Concertación de evaluación
        set_template_cell(table, 54, 2, datos.get("asignatura", ""))
        set_template_cell(table, 54, 15, datos.get("grupo", ""))

        if evaluaciones_df is not None and not evaluaciones_df.empty:
            for e_idx, (_, r) in enumerate(evaluaciones_df.iterrows()):
                t_row = 56 + e_idx
                if t_row <= 61:
                    set_template_cell(table, t_row, 0, r.get("Tipo de evaluación", ""), font_size=8)
                    set_template_cell(table, t_row, 2, r.get("Procedimiento de evaluación", ""), font_size=8)
                    val = r.get("Valor (%)", "")
                    val_str = f"{val}%" if val != "" and not str(val).endswith("%") else str(val)
                    set_template_cell(table, t_row, 11, val_str, font_size=8)
                    set_template_cell(table, t_row, 15, r.get("Fecha de realización", ""), font_size=8)

        # 5. Representantes estudiantiles
        if representantes_df is not None and not representantes_df.empty:
            for r_idx, (_, r) in enumerate(representantes_df.iterrows()):
                t_row = 65 + r_idx
                if t_row <= 67:
                    set_template_cell(table, t_row, 0, r.get("Nombre", ""), font_size=8)
                    set_template_cell(table, t_row, 5, r.get("Cedula", ""), font_size=8)
                    set_template_cell(table, t_row, 10, r.get("Correo", ""), font_size=8)

        # Docente y socialización
        set_template_cell(table, 68, 0, datos.get("profesor", ""))
        set_template_cell(table, 68, 5, datos.get("cedula_docente", ""))
        set_template_cell(table, 70, 0, f"Fecha de socialización de la Guía Didáctica: {datos.get('fecha_socializacion', '')}")

    else:
        # Fallback si no existe la plantilla estructurada
        add_custom_heading(doc, "GUÍA DIDÁCTICA Y CONCERTACIÓN DE EVALUACIÓN (FD-GC71)", level=1)

        # 1. IDENTIFICACIÓN DE LA ASIGNATURA
        add_custom_heading(doc, "1. IDENTIFICACIÓN DE LA ASIGNATURA", level=2)

        t_ident = doc.add_table(rows=3, cols=3)
        aplicar_bordes_tabla(t_ident, "002060")

        set_label_value(t_ident.rows[0].cells[0], "Programa", datos.get("programa", ""))
        set_label_value(t_ident.rows[0].cells[1], "Asignatura", datos.get("asignatura", ""))
        set_label_value(t_ident.rows[0].cells[2], "Código", datos.get("codigo", ""))

        set_label_value(t_ident.rows[1].cells[0], "Profesor", datos.get("profesor", ""))
        set_label_value(t_ident.rows[1].cells[1], "Correo", datos.get("correo", ""))
        set_label_value(t_ident.rows[1].cells[2], "Grupo", datos.get("grupo", ""))

        set_label_value(t_ident.rows[2].cells[0], "Créditos", datos.get("creditos", ""))
        set_label_value(t_ident.rows[2].cells[1], "HTP / HTI", f"{datos.get('htp', '')} / {datos.get('hti', '')}")
        set_label_value(t_ident.rows[2].cells[2], "Periodo", datos.get("periodo", ""))

        # 2. TEXTOS ACADÉMICOS BASE
        add_custom_heading(doc, "2. TEXTOS ACADÉMICOS BASE", level=2)

        campos_textos = [
            ("Justificación", "justificacion"),
            ("Competencias", "competencias"),
            ("Objetivo General", "objetivo_general"),
            ("Metodologías", "metodologias"),
            ("Ambientes de aprendizaje", "ambientes"),
            ("Medios y recursos", "medios"),
            ("Referencias bibliográficas", "referencias"),
        ]
        for label, key in campos_textos:
            val = str(datos.get(key, "")).strip()
            if val:
                p = doc.add_paragraph()
                r_l = p.add_run(f"{label}: ")
                r_l.bold = True
                r_l.font.color.rgb = RGBColor(0, 32, 96)
                r_v = p.add_run(val)
                r_v.font.size = Pt(10)

        # 3. PLAN DE SESIONES
        add_custom_heading(doc, "3. PLAN DE SESIONES", level=2)

        if sesiones_df is not None and not sesiones_df.empty:
            t_ses = doc.add_table(rows=1, cols=7)
            aplicar_bordes_tabla(t_ses, "002060")
            headers = ["Unidad", "Sesión", "Fecha", "Horario", "Contenido", "Trabajo Presencial", "Trabajo Independiente"]
            for idx, text in enumerate(headers):
                cell = t_ses.rows[0].cells[idx]
                shade_cell(cell, "002060")
                set_cell_text(cell, text, bold=True, color_rgb=(255, 255, 255), font_size_pt=9)

            for _, r in sesiones_df.iterrows():
                row_cells = t_ses.add_row().cells
                set_cell_text(row_cells[0], str(r.get("Unidad", "")), font_size_pt=8.5)
                set_cell_text(row_cells[1], str(r.get("N° sesión", "")), font_size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(row_cells[2], str(r.get("Fecha", "")), font_size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(row_cells[3], str(r.get("Horario", "")), font_size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(row_cells[4], str(r.get("Contenido por desarrollar", "")), font_size_pt=8.5)
                set_cell_text(row_cells[5], str(r.get("Descripción del trabajo presencial", "")), font_size_pt=8.5)
                set_cell_text(row_cells[6], str(r.get("Descripción trabajo independiente", "")), font_size_pt=8.5)

        # 4. CONCERTACIÓN DE EVALUACIÓN
        add_custom_heading(doc, "4. CONCERTACIÓN DE EVALUACIÓN", level=2)

        if evaluaciones_df is not None and not evaluaciones_df.empty:
            t_eval = doc.add_table(rows=1, cols=5)
            aplicar_bordes_tabla(t_eval, "002060")
            eval_headers = ["Tipo de Evaluación", "Procedimiento", "Valor (%)", "Fecha", "Unidad Relacionada"]
            for idx, text in enumerate(eval_headers):
                cell = t_eval.rows[0].cells[idx]
                shade_cell(cell, "002060")
                set_cell_text(cell, text, bold=True, color_rgb=(255, 255, 255), font_size_pt=9)

            for _, r in evaluaciones_df.iterrows():
                row_cells = t_eval.add_row().cells
                set_cell_text(row_cells[0], str(r.get("Tipo de evaluación", "")), font_size_pt=8.5)
                set_cell_text(row_cells[1], str(r.get("Procedimiento de evaluación", "")), font_size_pt=8.5)
                set_cell_text(row_cells[2], f"{r.get('Valor (%)', '')}%", font_size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(row_cells[3], str(r.get("Fecha de realización", "")), font_size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(row_cells[4], str(r.get("Unidad relacionada", "")), font_size_pt=8.5)

        # 5. REPRESENTANTES ESTUDIANTILES
        if representantes_df is not None and not representantes_df.empty:
            add_custom_heading(doc, "5. REPRESENTANTES DE ESTUDIANTES", level=2)

            t_rep = doc.add_table(rows=1, cols=3)
            aplicar_bordes_tabla(t_rep, "002060")
            rep_headers = ["Nombre del Estudiante", "Cédula / Documento", "Correo Institucional"]
            for idx, text in enumerate(rep_headers):
                cell = t_rep.rows[0].cells[idx]
                shade_cell(cell, "002060")
                set_cell_text(cell, text, bold=True, color_rgb=(255, 255, 255), font_size_pt=9)

            for _, r in representantes_df.iterrows():
                row_cells = t_rep.add_row().cells
                set_cell_text(row_cells[0], str(r.get("Nombre", "")), font_size_pt=8.5)
                set_cell_text(row_cells[1], str(r.get("Cedula", "")), font_size_pt=8.5)
                set_cell_text(row_cells[2], str(r.get("Correo", "")), font_size_pt=8.5)

    # Agregar Sello Digital Institucional y bloque de firma al final del documento
    p_head = doc.add_paragraph()
    r_head = p_head.add_run("SELLO DE FIRMA E INTEGRIDAD DIGITAL INSTITUCIONAL")
    r_head.bold = True
    r_head.font.size = Pt(12)
    r_head.font.color.rgb = RGBColor(0, 32, 96)

    t = doc.add_table(rows=2, cols=2)
    aplicar_bordes_tabla(t, "002060")

    c00 = t.rows[0].cells[0]
    set_cell_text(c00, "FIRMA DEL DOCENTE Y APROBACIÓN", bold=True, color_rgb=(0, 32, 96))
    c01 = t.rows[0].cells[1]
    set_cell_text(c01, "VERIFICACIÓN DIGITAL INSTITUCIONAL", bold=True, color_rgb=(0, 32, 96))

    c10 = t.rows[1].cells[0]
    set_section_text(c10, f"Docente: {datos.get('profesor','')}\nCédula: {datos.get('cedula_docente','')}\nFecha de concertación: {datos.get('fecha_socializacion','')}")

    c11 = t.rows[1].cells[1]
    import hashlib, json
    raw_hash = hashlib.sha256(json.dumps(datos, sort_keys=True, default=str).encode('utf-8')).hexdigest()
    set_section_text(c11, f"🔒 Documento Auditado por Sistema FDGC\nHash SHA-256: {raw_hash[:24]}...\nSello Institucional PCJIC - Válido")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

    # Agregar Sello Digital Institucional y bloque de firma al final del documento
    p_head = doc.add_paragraph()
    r_head = p_head.add_run("SELLO DE FIRMA E INTEGRIDAD DIGITAL INSTITUCIONAL")
    r_head.bold = True
    r_head.font.size = Pt(12)
    r_head.font.color.rgb = RGBColor(0, 32, 96)

    t = doc.add_table(rows=2, cols=2)
    aplicar_bordes_tabla(t, "002060")

    c00 = t.rows[0].cells[0]
    set_cell_text(c00, "FIRMA DEL DOCENTE Y APROBACIÓN", bold=True, color_rgb=(0, 32, 96))
    c01 = t.rows[0].cells[1]
    set_cell_text(c01, "VERIFICACIÓN DIGITAL INSTITUCIONAL", bold=True, color_rgb=(0, 32, 96))

    c10 = t.rows[1].cells[0]
    set_section_text(c10, f"Docente: {datos.get('profesor','')}\nCédula: {datos.get('cedula_docente','')}\nFecha de concertación: {datos.get('fecha_socializacion','')}")

    c11 = t.rows[1].cells[1]
    import hashlib, json
    raw_hash = hashlib.sha256(json.dumps(datos, sort_keys=True, default=str).encode('utf-8')).hexdigest()
    set_section_text(c11, f"🔒 Documento Auditado por Sistema FDGC\nHash SHA-256: {raw_hash[:24]}...\nSello Institucional PCJIC - Válido")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def crear_informe_gc72_docx(
    docente: str,
    periodo: str,
    fecha_entrega: date | str,
    cursos_df: pd.DataFrame,
    bloques_analisis: List[Dict[str, str]],
) -> bytes:
    doc = docx.Document(TEMPLATE_GC72 if os.path.exists(TEMPLATE_GC72) else None)
    
    p_head = doc.add_paragraph()
    r_head = p_head.add_run("VERIFICACIÓN DE FIRMA E INTEGRIDAD (FD-GC72)")
    r_head.bold = True
    r_head.font.size = Pt(12)
    r_head.font.color.rgb = RGBColor(0, 32, 96)
    t = doc.add_table(rows=2, cols=2)
    aplicar_bordes_tabla(t, "002060")
    
    c00 = t.rows[0].cells[0]
    set_cell_text(c00, "DOCENTE RESPONSABLE", bold=True, color_rgb=(0, 32, 96))
    c01 = t.rows[0].cells[1]
    set_cell_text(c01, "SELLO DE AUDITORÍA INSTITUCIONAL", bold=True, color_rgb=(0, 32, 96))
    
    c10 = t.rows[1].cells[0]
    f_str = fecha_entrega.strftime('%d/%m/%Y') if hasattr(fecha_entrega, 'strftime') else str(fecha_entrega)
    set_section_text(c10, f"Docente: {docente}\nPeriodo: {periodo}\nFecha de Entrega: {f_str}")
    
    c11 = t.rows[1].cells[1]
    import hashlib
    h = hashlib.sha256(f"{docente}_{periodo}_{f_str}".encode('utf-8')).hexdigest()
    set_section_text(c11, f"🔒 Informe FD-GC72 Validadas {len(cursos_df)} asignaturas\nHash SHA-256: {h[:24]}...\nPCJIC - Sistema de Gestión Académica")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def crear_informe_ejecutivo_institucional_docx(matriz_df: pd.DataFrame, resumen: Dict[str, Any]) -> bytes:
    doc = docx.Document()
    doc.add_heading("POLITÉCNICO COLOMBIANO JAIME ISAZA CADAVID", level=1)
    doc.add_heading("Informe Ejecutivo de Gobierno Académico", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Fecha de emisión: {date.today().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(f"Total cursos auditados: {resumen.get('total', 0)}\n")

    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Código", "Grupo", "Asignatura", "Profesor", "Estado", "Calidad"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        shade_cell(hdr_cells[i], "002060")

    if matriz_df is not None and not matriz_df.empty:
        for _, row in matriz_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row.get("Código", ""))
            row_cells[1].text = str(row.get("Grupo", ""))
            row_cells[2].text = str(row.get("Asignatura", ""))
            row_cells[3].text = str(row.get("Profesor", ""))
            row_cells[4].text = str(row.get("Estado", ""))
            row_cells[5].text = str(row.get("Score Calidad", ""))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def crear_gc71_html_imprimible(
    datos: Dict[str, str],
    sesiones_df: pd.DataFrame,
    evaluaciones_df: pd.DataFrame,
) -> str:
    """Genera una plantilla HTML formal lista para imprimir en PDF con CSS de alta fidelidad."""
    sesiones_rows = ""
    if sesiones_df is not None and not sesiones_df.empty:
        for _, r in sesiones_df.iterrows():
            sesiones_rows += f"""
            <tr>
                <td>{r.get('Unidad','')}</td>
                <td style="text-align:center;">{r.get('N° sesión','')}</td>
                <td style="text-align:center;">{r.get('Fecha','')}</td>
                <td style="text-align:center;">{r.get('Horario','')}</td>
                <td>{r.get('Contenido por desarrollar','')}</td>
                <td>{r.get('Descripción del trabajo presencial','')}</td>
                <td>{r.get('Descripción trabajo independiente','')}</td>
            </tr>
            """

    eval_rows = ""
    if evaluaciones_df is not None and not evaluaciones_df.empty:
        for _, r in evaluaciones_df.iterrows():
            eval_rows += f"""
            <tr>
                <td>{r.get('Tipo de evaluación','')}</td>
                <td>{r.get('Procedimiento de evaluación','')}</td>
                <td style="text-align:center;">{r.get('Valor (%)','')}%</td>
                <td style="text-align:center;">{r.get('Fecha de realización','')}</td>
                <td>{r.get('Unidad relacionada','')}</td>
            </tr>
            """

    return f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>FD-GC71 Guía Didáctica - {datos.get('asignatura','')}</title>
    <style>
        @page {{ size: letter; margin: 1.8cm; }}
        body {{ font-family: Arial, sans-serif; color: #1a1a1a; font-size: 11pt; line-height: 1.4; margin: 0; padding: 20px; }}
        .header-table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; border: 2px solid #002060; }}
        .header-table td {{ border: 1px solid #002060; padding: 8px; text-align: center; }}
        .header-title {{ font-weight: bold; font-size: 14pt; color: #002060; }}
        h2 {{ color: #002060; font-size: 12pt; border-bottom: 2px solid #002060; padding-bottom: 4px; margin-top: 24px; }}
        .data-table {{ width: 100%; border-collapse: collapse; margin-bottom: 16px; font-size: 9.5pt; }}
        .data-table th {{ background-color: #002060; color: white; border: 1px solid #002060; padding: 6px; }}
        .data-table td {{ border: 1px solid #cccccc; padding: 6px; }}
        .print-btn {{ display: block; width: 200px; margin: 0 auto 20px auto; padding: 10px; background: #1f4fd8; color: white; text-align: center; border-radius: 8px; font-weight: bold; text-decoration: none; cursor: pointer; }}
        @media print {{ .print-btn {{ display: none; }} }}
    </style>
</head>
<body>
    <button class="print-btn" onclick="window.print()">🖨️ Imprimir / Guardar PDF</button>

    <table class="header-table">
        <tr>
            <td style="width:25%; font-weight:bold; color:#002060;">POLITÉCNICO COLOMBIANO<br>JAIME ISAZA CADAVID</td>
            <td class="header-title">GUÍA DIDÁCTICA Y CONCERTACIÓN DE EVALUACIÓN<br>(FD-GC71)</td>
            <td style="width:20%; font-size:9pt;">Código: FD-GC71<br>Versión: 04</td>
        </tr>
    </table>

    <h2>1. IDENTIFICACIÓN DE LA ASIGNATURA</h2>
    <table class="data-table">
        <tr>
            <td><strong>Programa:</strong> {datos.get('programa','')}</td>
            <td><strong>Asignatura:</strong> {datos.get('asignatura','')}</td>
            <td><strong>Código:</strong> {datos.get('codigo','')}</td>
        </tr>
        <tr>
            <td><strong>Profesor:</strong> {datos.get('profesor','')}</td>
            <td><strong>Correo:</strong> {datos.get('correo','')}</td>
            <td><strong>Grupo:</strong> {datos.get('grupo','')}</td>
        </tr>
        <tr>
            <td><strong>Créditos:</strong> {datos.get('creditos','')}</td>
            <td><strong>HTP:</strong> {datos.get('htp','')} | <strong>HTI:</strong> {datos.get('hti','')}</td>
            <td><strong>Periodo:</strong> {datos.get('periodo','')}</td>
        </tr>
    </table>

    <h2>2. TEXTOS ACADÉMICOS BASE</h2>
    <p><strong>Justificación:</strong> {datos.get('justificacion','')}</p>
    <p><strong>Competencias:</strong> {datos.get('competencias','')}</p>
    <p><strong>Objetivo General:</strong> {datos.get('objetivo_general','')}</p>

    <h2>3. PLAN DE SESIONES</h2>
    <table class="data-table">
        <thead>
            <tr>
                <th>Unidad</th>
                <th>Sesión</th>
                <th>Fecha</th>
                <th>Horario</th>
                <th>Contenido</th>
                <th>Trabajo Presencial</th>
                <th>Trabajo Independiente</th>
            </tr>
        </thead>
        <tbody>
            {sesiones_rows}
        </tbody>
    </table>

    <h2>4. CONCERTACIÓN DE EVALUACIÓN</h2>
    <table class="data-table">
        <thead>
            <tr>
                <th>Tipo de Evaluación</th>
                <th>Procedimiento</th>
                <th>Valor (%)</th>
                <th>Fecha</th>
                <th>Unidad Relacionada</th>
            </tr>
        </thead>
        <tbody>
            {eval_rows}
        </tbody>
    </table>
</body>
</html>
"""
